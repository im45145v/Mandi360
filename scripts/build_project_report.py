"""Build the Mandi360 academic project report from verified local artifacts.

The script reads only already-produced project results and screenshot assets.
It never modifies the raw Google Maps exports.  The resulting DOCX is the
source for the submission PDF.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results" / "tables"
ASSETS = ROOT / "results" / "report_assets_v2"
OUTPUT = ROOT / "results" / "Mandi360_WAI_Project_Report_Malla_Venkata_Sai_Ashish_B31-25.docx"

REPORT_DATE = "16 August 2026"
TITLE = "Mandi @ 36: Customer Experience Intelligence"
SUBTITLE = "A Multi-Branch Data Mining and Predictive Analytics Study of a Hyderabad Restaurant Brand"
INK = "6F1D2A"
GOLD = "B9912F"
PALE_GOLD = "F7E9C4"
PALE_ROSE = "F8F1F2"
PALE_GRAY = "F1F3F5"


def read_json(relative: str) -> dict:
    return json.loads((RESULTS / relative).read_text(encoding="utf-8"))


def read_csv(relative: str) -> list[dict[str, str]]:
    with (RESULTS / relative).open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths_dxa = [round(width * 1440) for width in widths]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid_columns = list(table._tbl.tblGrid.gridCol_lst)
    if len(grid_columns) != len(widths):
        raise ValueError("Table grid column count does not match requested widths")
    for grid_column, width_dxa in zip(grid_columns, widths_dxa):
        grid_column.set(qn("w:w"), str(width_dxa))
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    tr_pr.append(marker)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_bottom_border(paragraph, color: str = GOLD, size: str = "10", space: str = "6") -> None:
    """Add a restrained accent rule below a major report heading."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_run_font(run, size: float = 12, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(r_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.append(size)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Malla Venkata Sai Ashish | B31-25 | Page ")
    set_run_font(run, size=9, color="666666")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    for name, size in (("Heading 1", 12), ("Heading 2", 12), ("Heading 3", 12)):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.keep_with_next = True

    caption = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_after = Pt(6)

    code = document.styles.add_style("Report Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code.font.size = Pt(9)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("Mandi @ 36 | WAI Project Report")
    set_run_font(header_run, size=9, color="666666")
    add_bottom_border(header, color=GOLD, size="6", space="4")

    list_bullet = document.styles["List Bullet"]
    list_bullet.font.name = "Times New Roman"
    list_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    list_bullet.font.size = Pt(12)
    list_bullet.paragraph_format.line_spacing = 1.5
    list_bullet.paragraph_format.space_after = Pt(0)


def paragraph(document: Document, text: str = "", *, style: str | None = None, align=None, before: float = 0, after: float = 0, keep: bool = False):
    p = document.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if keep:
        set_keep_with_next(p)
    return p


def heading(document: Document, text: str, level: int = 1):
    p = document.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, 12, bold=True)
    p.paragraph_format.line_spacing = 1.5
    if level == 1:
        add_bottom_border(p)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor.from_string(INK)
    return p


def bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r)


def add_caption(document: Document, number: int, text: str, source: str | None = None) -> None:
    p = document.add_paragraph(style="Figure Caption")
    r = p.add_run(f"Figure {number}. {text}")
    set_run_font(r, size=10, bold=True)
    if source:
        p2 = document.add_paragraph(style="Figure Caption")
        r2 = p2.add_run(source)
        set_run_font(r2, size=9, italic=True, color="555555")


def add_image(document: Document, filename: str, number: int, caption: str, width: float = 6.2) -> None:
    document.add_picture(str(ASSETS / filename), width=Inches(width))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(document, number, caption, "Source: Mandi @ 36 deployed dashboard, captured 16 August 2026.")


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float], note: str | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for cell, header in zip(header_cells, headers):
        set_cell_shading(cell, INK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=10, bold=True, color="FFFFFF")
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for cell, value in zip(cells, row):
            set_cell_shading(cell, PALE_GRAY if index % 2 else "FFFFFF")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=10)
    set_table_widths(table, widths)
    if note:
        p = paragraph(document, note, align=WD_ALIGN_PARAGRAPH.LEFT, before=3, after=6)
        for r in p.runs:
            set_run_font(r, size=9, italic=True, color="555555")


def add_cover(document: Document) -> None:
    for _ in range(3):
        paragraph(document, "", after=0)
    document.add_picture(str(ASSETS / "iim_ranchi_logo.png"), width=Inches(1.0))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = paragraph(document, "MANDI @ 36", align=WD_ALIGN_PARAGRAPH.CENTER, keep=True)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        set_run_font(r, size=24, bold=True, color=INK)
    p = paragraph(document, "Customer Experience Intelligence", align=WD_ALIGN_PARAGRAPH.CENTER, keep=True)
    p.paragraph_format.space_after = Pt(10)
    for r in p.runs:
        set_run_font(r, size=14, italic=True, color=GOLD)
    p = paragraph(document, SUBTITLE, align=WD_ALIGN_PARAGRAPH.CENTER, keep=True, after=18)
    for r in p.runs:
        set_run_font(r, size=12, italic=True)
    cover_rule = paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_bottom_border(cover_rule, color=GOLD, size="14", space="1")
    add_table(document,
        ["Report details", "Submission information"],
        [
            ["Course", "Data Mining and Predictive Analytics | Term IV | AY 2026-27"],
            ["Report type", "Individual Working with AI Project Report"],
            ["Submitted by", "Malla Venkata Sai Ashish"],
            ["Roll number", "B31-25"],
            ["Faculty", "Prof. Pradip Kumar Bala"],
            ["Submission date", REPORT_DATE],
        ],
        [2.0, 4.3],
    )
    paragraph(document, "", after=10)
    p = paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("Live deployment, click to open: ")
    set_run_font(r, size=12, bold=True)
    add_hyperlink(p, "https://mandi36.streamlit.app/", "https://mandi36.streamlit.app/")
    p2 = paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    r2 = p2.add_run("Source repository, click to open: ")
    set_run_font(r2, size=12, bold=True)
    add_hyperlink(p2, "https://github.com/im45145v/Mandi360", "https://github.com/im45145v/Mandi360")
    document.add_page_break()


def add_contents(document: Document) -> None:
    heading(document, "Report Guide")
    paragraph(document, "This report is organized around the six required Working with AI components. The sections below are supported by live dashboard evidence, two recorded AI Analyst scenarios, and referenced source links.")
    add_table(document,
        ["Section", "Focus"],
        [
            ["Executive Summary", "Project scope, verified dataset scale, and dashboard access."],
            ["(i) Business Objectives", "Decision problem, rationale, and review-source provenance."],
            ["(ii) Methodology Used", "Reproducible data pipeline, governance, and analytical methods."],
            ["(iii) Details of AI Use", "Guardrailed evidence agent and safeguards."],
            ["(iv) Prompt Use and Output", "Two live scenario questions, routing, evidence, and outputs."],
            ["(v) Results and Discussion", "Observed metrics, text mining, forecasting, and anomaly results."],
            ["(vi) Business Implications", "Operational routine, recommendations, and limitations."],
            ["Appendices and References", "Dashboard access, reproducibility resources, and APA references."],
        ],
        [2.15, 4.15],
    )
    document.add_page_break()


def build_report() -> None:
    summary = read_json("dataset_summary.json")
    branches = read_csv("branch_summary.csv")
    forecasts = read_json("predictive/branch_forecasts.json")
    clusters = read_json("clustering/cluster_model_kmeans.json")
    topics = read_json("nlp/topic_model_nmf.json")
    anomalies = read_json("anomaly/anomalies.json")
    document = Document()
    configure_styles(document)
    add_cover(document)
    add_contents(document)

    heading(document, "Executive Summary")
    p = paragraph(document, (
        "Mandi @ 36 is a customer experience intelligence system developed for a Hyderabad restaurant brand. "
        "The project converts collected Google Maps reviews into a manager-friendly Streamlit dashboard that supports "
        "data mining, predictive analytics, and evidence-led branch review. It combines review-volume and "
        "rating analysis with text-based sentiment, topic discovery, clustering, association rules, anomaly detection, "
        "linear trend forecasting, and an evidence-grounded AI Analyst. The dashboard is available for assessment at the "
    ))
    add_hyperlink(p, "https://mandi36.streamlit.app/", "https://mandi36.streamlit.app/")
    p.add_run(" and the complete reproducible codebase is available in the ")
    add_hyperlink(p, "https://github.com/im45145v/Mandi360", "https://github.com/im45145v/Mandi360")
    p.add_run(".")
    paragraph(document, (
        f"The analytical dataset contains {summary['record_count']:,} collected reviews across Banjara Hills, Gachibowli, and Jubilee Hills. "
        f"The overall observed average rating is {summary['average_rating']:.4f} on a five-point scale. Review text is available for "
        f"{summary['text_available_count']:,} records, so text-based findings are explicitly limited to that subset. The project deliberately "
        "separates collected facts, derived model outputs, and predicted outputs. This design makes the tool useful for management review without "
        "presenting correlation or forecasts as certainty."
    ))
    add_table(document,
        ["Indicator", "Verified value", "Evidence type"],
        [
            ["Collected review records", f"{summary['record_count']:,}", "Collected"],
            ["Branches", str(summary['branch_count']), "Collected"],
            ["Observed average rating", f"{summary['average_rating']:.4f} / 5", "Collected"],
            ["Reviews with usable text", f"{summary['text_available_count']:,}", "Collected"],
        ],
        [2.7, 1.8, 2.0],
        "Note. All values in this table are computed from collected review exports after schema validation."
    )
    paragraph(document, "The report follows the six required project-report parts: business objectives, methodology, AI use, prompt and output, results and discussion, and business implications.")

    heading(document, "(i) Business Objectives with Rationale Behind the Problem")
    paragraph(document, (
        "Restaurant managers receive a large volume of unstructured customer feedback but often lack a systematic way to compare branches, "
        "detect deteriorating experience, or translate patterns into timely actions. Manual reading is slow, inconsistent, and difficult to scale "
        "across locations. Mandi @ 36 addresses this problem by creating one evidence-oriented workspace for the brand and its three Hyderabad branches."
    ))
    heading(document, "Business Objectives", 2)
    bullet(document, "Consolidate customer feedback from the three branches into one validated and normalized review table.")
    bullet(document, "Compare observed rating, review volume, text availability, and experience signals by branch.")
    bullet(document, "Discover recurring language patterns, review groups, and issue combinations that warrant manager investigation.")
    bullet(document, "Identify unusual branch-month patterns and use transparent forecasts as an early-warning signal.")
    bullet(document, "Convert analytical evidence into a branch action-priority view, including issue priorities, suggested ownership, and review cycles.")
    bullet(document, "Allow a manager to ask natural-language questions while keeping the LLM constrained to compact, evidence-based inputs.")
    heading(document, "Why This Matters", 2)
    paragraph(document, (
        "The operational value lies in shortening the path from customer feedback to action. Instead of navigating individual review pages, a manager can "
        "see branch performance, drill into experience areas, review alert months, and compare predicted direction with historical evidence. This aligns "
        "with the course focus on applying data mining and predictive analytics to a real business problem and interpreting outputs for managerial decision-making "
        "(Larose & Larose, 2015)."
    ))
    add_table(document,
        ["Branch", "Collected reviews", "Observed average rating", "Reviews with text"],
        [[b["branch_name"], f"{int(b['review_count']):,}", f"{float(b['average_rating']):.4f}", f"{int(b['text_available_count']):,}"] for b in branches],
        [1.55, 1.45, 1.85, 1.65],
        "Note. Ratings and counts are collected-data metrics. The different branch volumes are a key interpretation limitation."
    )
    heading(document, "Scope and Source Locations", 2)
    p = paragraph(document, "The data were collected from Google Maps review pages using Apify's Google Maps scraping workflow. The supplied location pins were verified against the visible Google Maps address fields. The following branch locations provide source provenance and can be opened for context: ", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_hyperlink(p, "Banjara Hills", "https://maps.app.goo.gl/Q3ThcUTcjc2UsFoS8")
    p.add_run("; ")
    add_hyperlink(p, "Gachibowli", "https://maps.app.goo.gl/5rGRF4CdBzZYRULw8")
    p.add_run("; and ")
    add_hyperlink(p, "Jubilee Hills", "https://maps.app.goo.gl/DTt6HFX2THJhiixg8")
    p.add_run(".")
    for r in p.runs:
        if r.font.size is None:
            set_run_font(r)
    for branch_label, maps_url in [
        ("Banjara Hills, Alcazar Plaza, Road No. 1, click to open: ", "https://maps.app.goo.gl/Q3ThcUTcjc2UsFoS8"),
        ("Gachibowli, Lumbini Avenue, click to open: ", "https://maps.app.goo.gl/5rGRF4CdBzZYRULw8"),
        ("Jubilee Hills, Road No. 36, click to open: ", "https://maps.app.goo.gl/DTt6HFX2THJhiixg8"),
    ]:
        source_p = paragraph(document, "", align=WD_ALIGN_PARAGRAPH.LEFT)
        source_run = source_p.add_run(branch_label)
        set_run_font(source_run, size=10, bold=True)
        add_hyperlink(source_p, maps_url, maps_url)

    heading(document, "(ii) Methodology Used")
    paragraph(document, (
        "The methodology was designed as a reproducible Python pipeline. Raw branch exports are read only. The pipeline first normalizes schema differences, validates data quality, creates a privacy-minimized analytical table, writes deterministic analytical artifacts, and then presents those artifacts in the dashboard. "
        "The LLM layer is deliberately separated from deterministic analytics."
    ))
    add_table(document,
        ["Stage", "What happens", "Control"],
        [
            ["1. Collection", "Apify exports Google Maps reviews for three branch locations.", "Source provenance retained."],
            ["2. Data foundation", "Schema detection, normalization, and validation produce an analytical table.", "Raw exports are not modified."],
            ["3. Analytics", "EDA, NLP, topics, clustering, association rules, anomaly detection, forecasting, and action-priority signals are produced.", "Results are tagged collected, derived, or predicted."],
            ["4. Decision support", "The Streamlit dashboard exposes separate management views and the AI Analyst uses compact evidence bundles.", "LLM does not receive the full dataset."],
        ],
        [1.3, 3.1, 2.1],
        "Note. The methodology intentionally keeps deterministic analytics separate from the LLM feature."
    )
    heading(document, "Data Preparation and Governance", 2)
    paragraph(document, (
        "The three Apify-style JSON exports are not structurally identical. Banjara Hills and Gachibowli use flatter records, while Jubilee Hills contains nested reviewer and place structures. The ingestion module detects the record shape and maps it into a common schema containing review identifier, branch, date, rating, review text, language, owner response, and provenance. Validation checks inspect missing values, rating range, duplicate identifiers, generated identifiers, branch counts, and the observed date range. Reviewer profile URLs, photographs, and other unnecessary scraper metadata are excluded from the analytical output."
    ))
    heading(document, "Analytical Methods", 2)
    add_table(document,
        ["Component", "Method", "Purpose and interpretation"],
        [
            ["Descriptive analysis", "Branch and monthly aggregation", "Observed ratings, review counts, and coverage; collected evidence."],
            ["Sentiment and aspects", "Baseline lexicon and keyword aspects", "Derived customer-mood and experience-area signals; review examples should be checked."],
            ["Topic discovery", "TF-IDF plus NMF", "Eight derived text themes from reviews with usable text."],
            ["Review clustering", "TF-IDF plus KMeans; silhouette selection", "Eight recurring review groups; group labels require human interpretation."],
            ["Association rules", "Apriori, support 0.01, confidence 0.30", "Issue co-occurrence patterns, not causal relationships."],
            ["Anomaly detection", "Isolation Forest; contamination 0.10", "Unusual branch-month combinations for manager review, not confirmed incidents."],
            ["Forecasting", "Per-branch linear regression", "Three-month rating outlook and risk label; a transparent, unvalidated early-warning signal."],
            ["Action prioritization", "Evidence-ranked rules", "Priority issue, owner, and weekly review cycle."],
        ],
        [1.25, 1.75, 3.5],
        "Note. Derived and predicted methods are not presented as causal proof or certainty."
    )
    heading(document, "Dashboard Experience", 2)
    paragraph(document, (
        "The Streamlit application organizes the evidence into distinct pages: Executive Overview, Sentiment, Topics and Aspects, Data Mining, Predictive Analytics, Anomaly Center, Branch Intelligence, AI Analyst, and Negative Insights. This multipage design lets a user start with a brand-level question and then move to the relevant analytic layer. Streamlit supports interactive data applications with charts, tables, widgets, and multipage workflows (Streamlit, n.d.)."
    ))
    add_image(document, "01_overview.png", 1, "Executive Overview page, showing the full navigation, collected-data framing, and brand-level indicators.")
    add_image(document, "04_sentiment.png", 2, "Sentiment page, showing the branch selector, customer-mood mix, and branch-level mood chart.")
    add_image(document, "10_sentiment_wordcloud.png", 3, "Sentiment page word cloud: term size represents frequency in collected review text, while the chart below shows monthly derived customer-mood signals.")

    heading(document, "(iii) Details of AI Use")
    paragraph(document, (
        "AI is used only in the AI Analyst feature. The deterministic pipeline calculates and stores the analytical outputs first. When a manager asks a question, rule-based code identifies the branch, comparison intent, and any explicitly requested month. It then assembles a compact JSON evidence bundle. The bundle can contain the brand summary, branch performance, monthly trend, aspect sentiment, anomaly flags, forecast, topics, association rules, and action-priority signals. The full raw review dataset is not sent to the LLM."
    ))
    paragraph(document, (
        "When an API key is configured, the application uses the OpenAI chat-completions client with the gpt-4o-mini model. The system instruction requires the model to distinguish REAL values from DERIVED outputs and PREDICTED outlooks, avoid inventing facts, avoid causal claims from correlation, and state when evidence is sparse. The action-recommendation stage then asks for concrete prioritized actions, an evidence basis, and a confidence label. If the LLM service is not configured, the dashboard still presents the deterministic evidence table rather than pretending to generate an answer. This fallback is important for auditability and robust use."
    ))
    heading(document, "AI Safeguards", 2)
    bullet(document, "Rule-based routing is deterministic; the LLM does not decide which branch or period is in scope.")
    bullet(document, "The LLM receives compact aggregated evidence rather than 33,275 raw records or reviewer-identifying metadata.")
    bullet(document, "System instructions prohibit fabricated facts and causal claims from correlations.")
    bullet(document, "The interface labels collected, derived, and predicted evidence so the user can judge confidence appropriately.")
    bullet(document, "Manager validation is required before operational action, especially for forecasts and derived text signals.")

    heading(document, "(iv) Step-by-Step Use of Prompt in AI and the Output")
    paragraph(document, (
        "The following illustrates the actual in-application workflow for the default management question used in the deployed AI Analyst. It is presented as a reproducible prompt pathway rather than as an uncontrolled general-purpose chat interaction."
    ))
    heading(document, "Step 1: Manager Question", 2)
    p = paragraph(document, "How is Banjara Hills performing and what actions should we take?", align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.left_indent = Inches(0.25)
    for r in p.runs:
        set_run_font(r, size=11, italic=True)
    heading(document, "Step 2: Deterministic Routing", 2)
    paragraph(document, "The orchestration layer detects the Banjara Hills branch and recognizes that the phrase 'what actions should we take' requests evidence-led action recommendations. It selects the relevant evidence functions before any LLM call is made.")
    heading(document, "Step 3: Compact Evidence Bundle", 2)
    paragraph(document, "The AI Analyst compiles a branch summary, observed monthly performance, aspect-level signals, alert months, forecast information, top topics, association rules, and the relevant action-priority signals. This is the only analytical material forwarded for reasoning.")
    heading(document, "Step 4: Guardrailed Prompt", 2)
    p = document.add_paragraph(style="Report Code")
    p.add_run("Role: data-mining investigation analyst for the Mandi @ 36 restaurant brand\n")
    p.add_run("Required behavior: distinguish REAL, DERIVED, and PREDICTED evidence; do not invent facts; do not claim causation from correlation; state sparse evidence.\n")
    p.add_run("User question: How is Banjara Hills performing and what actions should we take?\n")
    p.add_run("Evidence: compact JSON bundle generated from precomputed artifacts.")
    heading(document, "Step 5: Evidence-Grounded Output", 2)
    paragraph(document, (
        "The resulting investigation reports observed monthly ratings and review counts, identifies alert reasons, and proposes actions such as staff-training review, feedback monitoring, and food-quality checks. Figure 4 shows the linked branch-evidence view used to interpret these suggestions. The output must be treated as a managerial decision aid. The ratings and counts are observed evidence; the narrative, aspect signals, and forecasts have different confidence levels and should be verified before action."
    ))
    add_image(document, "08_branch_intelligence.png", 4, "Branch Intelligence page for Banjara Hills, combining collected, derived, and predicted evidence used in focused investigation.")
    heading(document, "Live AI Analyst Scenario 1: Branch Investigation", 2)
    p = paragraph(document, "Question submitted: How is Banjara Hills performing and what actions should we take?", align=WD_ALIGN_PARAGRAPH.LEFT)
    for r in p.runs:
        set_run_font(r, size=11, italic=True)
    paragraph(document, (
        "The deployed AI Analyst returned a Banjara Hills performance summary, bringing together observed branch-month ratings, review counts, alert severity, and a prioritized action plan. The response recommended investigating alert months, reviewing service quality, monitoring future performance, and using feedback follow-up to verify and improve the customer experience. Each recommendation was presented with an evidence basis and confidence label."
    ))
    add_image(document, "11_ai_scenario_banjara_full_page.png", 5, "Live AI Analyst Scenario 1: complete AI Analyst page, submitted Banjara Hills question, and generated investigation output.", width=6.25)
    heading(document, "Live AI Analyst Scenario 2: Comparative Priority", 2)
    p = paragraph(document, "Question submitted: Compare Banjara Hills and Gachibowli. Which branch needs priority attention and why?", align=WD_ALIGN_PARAGRAPH.LEFT)
    for r in p.runs:
        set_run_font(r, size=11, italic=True)
    paragraph(document, (
        "The second live run compared branch-level ratings, alert patterns, and outlook signals. The answer prioritized Banjara Hills for immediate review because the evidence bundle contained several high-severity alert months and larger rating instability. This is a decision-support recommendation, not a causal conclusion, and the manager is expected to validate it with current review samples before action."
    ))
    add_image(document, "12_ai_scenario_comparison_full_page.png", 6, "Live AI Analyst Scenario 2: complete AI Analyst page, branch-comparison question, and generated priority investigation.", width=6.25)

    heading(document, "(v) Results and Discussion")
    paragraph(document, (
        "This section distinguishes observed results from model-derived outputs. The dataset contains 33,275 collected reviews between 20 May 2017 and 15 August 2026. The total includes 16,457 five-star reviews, 9,888 four-star reviews, 3,369 three-star reviews, 1,093 two-star reviews, and 2,468 one-star reviews. Overall ratings are strong, but branch volume is highly uneven: Jubilee Hills accounts for the majority of records. Therefore, raw volume should not be interpreted as a direct measure of branch quality."
    ))
    heading(document, "Observed Branch Comparison", 2)
    paragraph(document, (
        "Jubilee Hills has the highest observed average rating (4.1283) and the largest review volume (28,259). Banjara Hills has an average rating of 3.9913 across 3,000 reviews. Gachibowli has the lowest observed average rating (3.9499) across 2,016 reviews. These differences are descriptive observations, not a causal evaluation of branch operations. The recommended interpretation is to use the comparison as a starting point for branch-level review."
    ))
    add_image(document, "02_branch_comparison_chart.png", 7, "Branch Comparison chart from the app: bar height represents observed average rating and the bar label gives the collected review count.")
    add_image(document, "03_rating_trend_chart.png", 8, "Observed monthly average-rating trend by branch, as displayed on the Executive Overview page.")
    heading(document, "Text Mining and Pattern Discovery", 2)
    paragraph(document, (
        f"Of the collected reviews, {summary['text_available_count']:,} have usable text. The topic model fitted {topics['params']['num_topics_fitted']} NMF topics using {topics['params']['documents_used']:,} documents and a vocabulary of {topics['params']['vocabulary_size']:,} terms. The review-clustering model selected {clusters['params']['selected_k']} groups from {clusters['params']['documents_used']:,} review texts. Common terms include food, taste, service, ambience, mutton, chicken, good, and best. These are language patterns rather than direct satisfaction scores and are intended to focus human reading of customer comments."
    ))
    add_image(document, "05_data_mining.png", 9, "Data Mining page, presenting the eight recurring customer-review groups and their distribution.")
    heading(document, "Early-Warning and Predictive Results", 2)
    flagged_count = sum(1 for row in anomalies["rows"] if row.get("is_anomaly"))
    paragraph(document, (
        f"The anomaly process analyzed {anomalies['params']['rows_used']} eligible branch-month observations using review count, average rating, average sentiment score, and negative ratio. It flagged {flagged_count} months for attention. An alert is not proof of an operational incident; it is a signal to inspect contemporary reviews and branch context. The dashboard makes this distinction explicit."
    ))
    forecast_rows = []
    for row in forecasts["branches_forecasted"]:
        evaluation = row["evaluation"]
        forecast_rows.append([
            row["branch_id"].replace("_", " ").title(),
            str(row["months_available"]),
            f"{row['trend_slope_per_month']:.6f}",
            row["risk_level"].title(),
            f"{evaluation['mae']:.4f}",
            "Yes" if evaluation["beats_naive_baseline"] else "No",
        ])
    add_table(document,
        ["Branch", "Months", "Trend per month", "Risk", "Holdout MAE", "Beats naive?"],
        forecast_rows,
        [1.25, 0.65, 1.15, 0.75, 1.05, 1.15],
        "Note. Forecasts are three-month linear-trend extrapolations and are marked predicted and unvalidated. Holdout MAE is included to make model quality visible."
    )
    paragraph(document, (
        "All three branches are currently classified as WATCH because their fitted monthly trend slopes are negative but not below the elevated-risk threshold. Banjara Hills is forecast at 3.7672 for September 2026, Gachibowli at 3.7128, and Jubilee Hills at 3.7335. These are model outputs, not actual future ratings. The Gachibowli and Banjara Hills models beat the naive holdout baseline, while the Jubilee Hills model did not. This is a useful caution against overconfident use of the forecast."
    ))
    add_image(document, "06_predictive_analytics.png", 10, "Predictive Analytics page: observed Banjara Hills ratings, transparent forecast, and likely range.")
    add_image(document, "07_anomaly_center.png", 11, "Anomaly Center: observed branch-month ratings with attention months visibly marked.")
    heading(document, "Evidence-Led Action Priorities", 2)
    paragraph(document, (
        "The deterministic action-priority table ranks recurring branch issues using derived evidence and supporting alert counts. For Banjara Hills, service, waiting time, food quality, ambience, hygiene, price/value, and quantity are high-priority review areas in the current artifact set. For example, the service signal has 298 mentions and seven supporting alert months, while food quality has 703 mentions and seven supporting alert months. These numbers indicate where managers should investigate, not that the source terms are proven causes of lower ratings."
    ))

    heading(document, "(vi) Business Implications")
    paragraph(document, (
        "Mandi @ 36 can improve customer-experience management by turning a large review archive into a disciplined weekly operating rhythm. At the brand level, leadership can compare observed performance across branches. At the branch level, managers can identify a specific experience area, read the supporting feedback samples, and assign an owner. At the early-warning level, the team can examine an unusual month before an issue becomes persistent. At the AI Analyst level, a decision-maker can ask a focused question and receive an answer grounded in the same prepared evidence used by the dashboard."
    ))
    heading(document, "Recommended Management Routine", 2)
    bullet(document, "Weekly branch huddle: review high-priority action signals, new alert months, and a small sample of customer comments.")
    bullet(document, "Monthly brand review: compare observed ratings and text coverage across branches, then ask why a difference may exist before changing operations.")
    bullet(document, "Service recovery: where service or waiting-time signals persist, review staffing, greeting, order follow-up, and peak-period process checks.")
    bullet(document, "Food and value review: where food quality, quantity, or price/value appears in the action-priority queue, inspect recipes, portions, menu communication, and recent low-rating comments.")
    bullet(document, "Forecast review: treat WATCH labels as an invitation to monitor leading indicators, not as a claim that customer satisfaction will decline.")
    bullet(document, "Governance: retain the REAL, DERIVED, and PREDICTED distinction in every managerial presentation, and use human validation for actions affecting staff, customers, or promotions.")
    heading(document, "Limitations", 2)
    paragraph(document, (
        "The findings represent available Google Maps reviews, not a random survey of all customers. Only 11,765 records contain review text, and the branch review volumes are markedly unequal. Sentiment, aspects, topics, clusters, association rules, and anomaly flags are model-derived signals. Association rules do not demonstrate cause and effect. The forecast model is intentionally simple and transparent, uses a short holdout evaluation, and remains unvalidated for production forecasting. These constraints are not defects to hide; they are essential conditions for responsible interpretation."
    ))
    paragraph(document, (
        "The project nevertheless demonstrates a strong end-to-end application of data mining and predictive analytics: real review collection, data normalization, descriptive analysis, unsupervised learning, anomaly detection, transparent forecasting, action prioritization, controlled LLM assistance, and a deployed application that makes the results accessible to non-technical users."
    ))

    heading(document, "Appendix A: Dashboard Access and Screens")
    p = paragraph(document, "Figures 1 to 11 show the app's different pages, charts, word-cloud section, and two complete AI Analyst scenarios rather than isolated navigation fragments. The user can test the complete application through the ")
    add_hyperlink(p, "https://mandi36.streamlit.app/", "https://mandi36.streamlit.app/")
    p.add_run(".")

    heading(document, "Appendix B: Reproducibility and Technical Resources")
    paragraph(document, (
        "The report is supported by an executable Python project. The ingestion module normalizes source records; preprocessing creates clean text fields; analytics modules generate EDA, NLP, topic, clustering, association, anomaly, forecast, and action-priority artifacts; the Streamlit app presents these artifacts. The complete source code and notebook are available in the project repository. The pipeline entry point is src/pipeline.py and the companion notebook is notebooks/data_mining_predictive_analytics.ipynb."
    ))
    p = paragraph(document, "", align=WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run("Repository, click to open: ")
    set_run_font(r, bold=True)
    add_hyperlink(p, "https://github.com/im45145v/Mandi360", "https://github.com/im45145v/Mandi360")
    p2 = paragraph(document, "", align=WD_ALIGN_PARAGRAPH.LEFT)
    r2 = p2.add_run("Live deployment, click to open: ")
    set_run_font(r2, bold=True)
    add_hyperlink(p2, "https://mandi36.streamlit.app/", "https://mandi36.streamlit.app/")
    paragraph(document, "Representative pipeline command:", keep=True)
    code = document.add_paragraph(style="Report Code")
    code.add_run("python -m src.pipeline\nstreamlit run app/Home.py")

    heading(document, "References")
    references = [
        "Apify. (n.d.). Google Maps Scraper. https://apify.com/apify/google-maps-scraper",
        "Google. (n.d.). Google Maps. Retrieved August 16, 2026, from https://www.google.com/maps",
        "Han, J., Kamber, M., & Pei, J. (2012). Data mining: Concepts and techniques (3rd ed.). Morgan Kaufmann.",
        "Larose, D. T., & Larose, C. D. (2015). Data mining and predictive analytics (2nd ed.). Wiley.",
        "OpenAI. (n.d.). OpenAI API documentation. https://platform.openai.com/docs/overview",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., VanderPlas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, E. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "Streamlit. (n.d.). Streamlit documentation. https://docs.streamlit.io/",
    ]
    for item in references:
        p = paragraph(document, item, align=WD_ALIGN_PARAGRAPH.LEFT, after=3)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    document.core_properties.title = TITLE
    document.core_properties.author = "Malla Venkata Sai Ashish"
    document.core_properties.subject = "Data Mining and Predictive Analytics WAI Project Report"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
