"""Create the CRM Track C report and prompt logbook from verified project artifacts.

The report uses only the already generated Google Maps review outputs and dashboard
screens. It never changes raw review exports.
"""
from __future__ import annotations

import csv
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "results" / "tables"
SCREENS = ROOT / "results" / "report_assets_v2"
ASSETS = ROOT / "results" / "crm_report_assets"
OUTPUT = ROOT / "results" / "Mandi360_CRM_TrackC_Report_Malla_Venkata_Sai_Ashish_B31-25.docx"
LOGBOOK_OUTPUT = ROOT / "results" / "Mandi360_CRM_Prompt_Logbook_Malla_Venkata_Sai_Ashish_B31-25.docx"

NAME = "Malla Venkata Sai Ashish"
ROLL = "B31-25"
REPORT_DATE = "17 August 2026"
INK = "6F1D2A"
GOLD = "B9912F"
PALE_GOLD = "F7E9C4"
PALE_ROSE = "F8F1F2"
PALE_GRAY = "F1F3F5"
BLUE = "1F4D78"
MUTED = "666666"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def set_run_font(run, size: float = 12, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    dxa = [round(width * 1440) for width in widths]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for grid, value in zip(table._tbl.tblGrid.gridCol_lst, dxa):
        grid.set(qn("w:w"), str(value))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    tr_pr.append(marker)


def bottom_rule(paragraph, color: str = GOLD) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relation = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    under = OxmlElement("w:u")
    under.set(qn("w:val"), "single")
    r_pr.append(under)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.append(size)
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{NAME} | {ROLL} | Page ")
    set_run_font(run, size=10, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run_el = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    props.append(fonts)
    run_el.append(props)
    text = OxmlElement("w:t")
    text.text = "1"
    run_el.append(text)
    field.append(run_el)
    paragraph._p.append(field)


def setup_document(document: Document, running_label: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(running_label)
    set_run_font(run, size=10, color=MUTED)
    bottom_rule(header, GOLD)
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in (
        ("Heading 1", 16, INK, 16, 8),
        ("Heading 2", 13, INK, 11, 5),
        ("Heading 3", 12, BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    if level == 1:
        bottom_rule(paragraph)


def add_para(document: Document, text: str = "", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before: float = 0, after: float = 7, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, italic=italic)


def add_labeled_para(document: Document, label: str, body: str, *, color: str = INK) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(5)
    lead = paragraph.add_run(label)
    set_run_font(lead, size=12, bold=True, color=color)
    detail = paragraph.add_run(body)
    set_run_font(detail, size=12)


def add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GOLD)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(1)
    label = paragraph.add_run(title + " ")
    set_run_font(label, size=11, bold=True, color=INK)
    detail = paragraph.add_run(body)
    set_run_font(detail, size=11)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    font_size: float = 10.5,
    left_aligned_cols: set[int] | None = None,
) -> None:
    if left_aligned_cols is None:
        left_aligned_cols = {0}
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        set_cell_shading(cell, INK)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=font_size, bold=True, color="FFFFFF")
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, (cell, value) in enumerate(zip(cells, values)):
            if row_index % 2:
                set_cell_shading(cell, PALE_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in left_aligned_cols else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image(document: Document, filename: str, width: float, caption: str, source: str | None = None) -> None:
    path = filename if filename.startswith("/") else str(SCREENS / filename)
    document.add_picture(path, width=Inches(width))
    inline = document.inline_shapes[-1]._inline
    inline.docPr.set("descr", caption)
    inline.docPr.set("title", "Mandi @ 36 CRM dashboard figure")
    picture = document.paragraphs[-1]
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p = document.add_paragraph(style="Caption")
    caption_p.add_run(caption).bold = True


def add_gantt_chart(document: Document) -> None:
    """Add an editable Word-native 12-week Gantt chart with exact table geometry."""
    tasks = [
        ("Define fields and ownership", 0, 2, INK),
        ("Pilot sentiment and alerts", 2, 3, GOLD),
        ("Train support playbooks", 4, 3, BLUE),
        ("Branch rollout and review", 7, 3, "376F77"),
        ("Evaluate and calibrate", 10, 2, "8B3A3A"),
    ]
    table = document.add_table(rows=1, cols=13)
    table.style = "Table Grid"
    set_table_widths(table, [2.42] + [0.34] * 12)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, cell in enumerate(header.cells):
        set_cell_shading(cell, INK)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("Activity" if index == 0 else f"W{index}")
        set_run_font(run, size=8.4, bold=True, color="FFFFFF")
    for label, start, duration, color in tasks:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        paragraph = cells[0].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        set_run_font(run, size=8.5, bold=True)
        for week in range(12):
            cell = cells[week + 1]
            if start <= week < start + duration:
                set_cell_shading(cell, color)
            else:
                set_cell_shading(cell, PALE_GRAY)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_cover(document: Document) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(8)
    document.add_picture(str(SCREENS / "iim_ranchi_logo.png"), width=Inches(1.0))
    logo = document.inline_shapes[-1]._inline
    logo.docPr.set("descr", "Indian Institute of Management Ranchi logo")
    logo.docPr.set("title", "IIM Ranchi logo")
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        document.add_paragraph().paragraph_format.space_after = Pt(8)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("CUSTOMER RELATIONSHIP MANAGEMENT PROJECT | TRACK C")
    set_run_font(run, size=11, bold=True, color=GOLD)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(9)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("AI-Driven Sentiment Analysis of Customer Feedback")
    set_run_font(run, size=24, bold=True, color=INK)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run("A CRM Study of Mandi @ 36, Hyderabad")
    set_run_font(run, size=15, italic=True, color=BLUE)
    table = document.add_table(rows=3, cols=2)
    set_table_widths(table, [2.0, 4.5])
    metadata = [("Submitted by", NAME), ("Roll Number", ROLL), ("Submission date", REPORT_DATE)]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], PALE_GOLD)
        for cell, text, bold in ((row.cells[0], label, True), (row.cells[1], value, False)):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, size=12, bold=bold, color=INK if bold else None)
    document.add_paragraph().paragraph_format.space_after = Pt(8)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Voice of Customer, sentiment intelligence, and evidence-led service recovery")
    set_run_font(run, size=11, italic=True, color=MUTED)
    document.add_page_break()


def build_report() -> None:
    dataset = read_json(TABLES / "dataset_summary.json")
    branches = read_csv(TABLES / "branch_summary.csv")
    crm_cases = read_csv(TABLES / "crm" / "crm_cases.csv")
    nlp_rows = read_csv(TABLES / "nlp" / "reviews_nlp_baseline.csv")
    aspects = read_csv(TABLES / "nlp" / "review_aspects_baseline.csv")
    topic_model = read_json(TABLES / "nlp" / "topic_model_nmf.json")
    anomalies = read_json(TABLES / "anomaly" / "anomalies.json")

    sentiment = Counter(row["sentiment_label"] for row in nlp_rows)
    aspect_counts = Counter(row["aspect"] for row in aspects)

    document = Document()
    setup_document(document, "Mandi @ 36 | CRM Track C Report")
    add_cover(document)

    add_heading(document, "Executive Summary")
    add_para(document, (
        "Mandi @ 36 receives customer feedback through Google Maps reviews across its Banjara Hills, "
        "Gachibowli, and Jubilee Hills branches. This CRM project converts that unstructured Voice of Customer "
        "material into a repeatable feedback-to-action workflow. It analyzes 33,275 collected reviews, of which "
        "11,765 contain usable text, using transparent sentiment, aspect, topic, alert, and prioritization methods."
    ))
    add_callout(document, "Core CRM insight.", (
        "The value is not simply a sentiment label. The dashboard combines review volume, experience-area mentions, "
        "rating trends, and alert signals so a branch manager can decide what to inspect, who should own it, and when "
        "the customer experience should be reviewed again."
    ))
    add_labeled_para(document, "Course alignment. ", (
        "The workflow applies the CRM course logic of customer-centricity, customer analytics, journey touchpoints, "
        "service response, and responsible AI. Google Maps reviews are treated as a post-experience touchpoint: they are "
        "not merely promotional content, but evidence that can be closed through a documented response and improvement cycle "
        "(Kumar & Reinartz, 2018; Winer, 2001)."
    ))
    add_labeled_para(document, "Finding snapshot. ", (
        f"Jubilee Hills has the highest observed average rating ({branches[2]['average_rating']}), while Gachibowli "
        f"has the lowest ({branches[1]['average_rating']}). The text pipeline finds food quality, ambience, service, "
        "waiting time, and price/value as recurring experience areas. These are analytical signals, not proof of causation."
    ))
    add_labeled_para(document, "CRM decision. ", (
        "Use negative-feedback alerts to trigger a short branch review, classify the issue, validate a sample of the "
        "underlying reviews, assign an owner, and log the service-recovery or process-improvement response."
    ))
    add_heading(document, "1. Problem Background and Identification")
    add_para(document, (
        "Restaurants receive feedback at a volume that makes manual reading inconsistent and slow. A rating alone "
        "cannot show whether dissatisfaction relates to food, service, wait time, value, or ambience. Without an "
        "organized Voice of Customer process, recurring pain points can remain unresolved, delayed responses can erode "
        "trust, and branch teams cannot consistently connect feedback to a CRM action."
    ))
    add_labeled_para(document, "Problem statement. ", (
        "Mandi @ 36 receives thousands of online customer reviews but needs a systematic, evidence-led way to detect "
        "recurring experience issues, prioritize response, and inform branch-level CRM action. This project applies "
        "AI-supported sentiment and text analysis to convert review text into monitored CRM signals."
    ))
    add_labeled_para(document, "Stakeholders. ", (
        "Branch managers, service teams, marketing and CRM coordinators, food operations, senior management, and "
        "customers who benefit when feedback is acknowledged and converted into visible improvement."
    ))

    add_heading(document, "2. Literature Review")
    add_para(document, (
        "Sentiment analysis is commonly used to identify opinions and evaluations expressed in text, while opinion "
        "targets and context determine whether the output is useful for management action (Liu, 2012; Pang & Lee, 2008). "
        "For customer feedback, aspect-based analysis is particularly important because a single restaurant review can "
        "praise food and criticize service in the same message."
    ))
    add_para(document, (
        "Topic modeling supports the discovery of recurring language patterns in large feedback collections, while current "
        "approaches such as BERTopic illustrate "
        "embedding-and-clustering alternatives (Grootendorst, 2022). In this project, a transparent TF-IDF and NMF method "
        "is used so terms can be inspected by a human before a topic is treated as a business theme."
    ))
    add_para(document, (
        "Voice of Customer research shows that text analytics can move feedback from broad sentiment toward service and "
        "product decisions. Enterprise VoC research expands beyond polarity to distinguish feedback, help-seeking, enquiry, "
        "and buying intent (Suresh et al., 2018). Hospitality studies have shown that online-review themes can inform service "
        "response and customer-experience design (Zhang, 2019; Özdağoğlu et al., 2018)."
    ))
    add_labeled_para(document, "Research gap addressed. ", (
        "The project does not claim that an automated label is a customer-satisfaction metric. It connects transparent NLP "
        "outputs to a documented CRM response routine, preserving review evidence for manager validation."
    ))

    add_heading(document, "3. Data Collection, Preparation, and Methodology")
    add_para(document, (
        "The project uses one feedback source: Google Maps customer reviews collected from three Mandi @ 36 locations "
        "through Apify exports. The available review window is 20 May 2017 to 15 August 2026. This full available window "
        "is disclosed because it is wider than a short pilot window and branch coverage is uneven. The raw files are preserved; "
        "the pipeline writes normalized and analytical outputs separately."
    ))
    add_table(document,
        ["Stage", "Method", "CRM use"],
        [
            ["Collect", "Apify Google Maps review exports", "Retain source, location, date, rating, and review text."],
            ["Prepare", "Schema detection, normalization, validation, text cleaning", "Create privacy-minimized analytical fields without editing raw exports."],
            ["Classify", "Transparent lexicon sentiment baseline", "Flag broad customer-mood signals for review."],
            ["Discover", "Keyword aspects and TF-IDF plus NMF topics", "Identify repeat experience areas and language themes."],
            ["Monitor", "Monthly trends and Isolation Forest alerts", "Surface unusual branch-month patterns for manager review."],
            ["Act", "Evidence-ranked CRM cases and AI-supported explanation", "Assign owner, action, and review cycle."],
        ], [1.05, 2.35, 3.10], font_size=9.5)
    add_para(document, "Note. Sentiment, aspects, themes, and alerts are derived or model-based signals. Ratings and review counts are collected evidence.", italic=True, after=4)
    add_labeled_para(document, "Customer 360 principle. ", (
        "The present dashboard provides a branch-level experience view by combining review text, rating, date, location, "
        "theme, alert, and CRM-case fields. A future customer-level Customer 360 view should only be built with consented "
        "identifiers, clear data purpose, access control, and a retention policy."
    ))
    add_heading(document, "4. Dashboard and AI Tools")
    add_para(document, (
        "The dashboard was built in Streamlit and separates collected, derived, and predicted outputs. The Sentiment page offers "
        "mood distribution, branch comparison, a word cloud, and a monthly trend. The Negative Insights page isolates complaint "
        "signals and suggested CRM fixes. The AI Analyst is a controlled assistant: deterministic code identifies the relevant branch "
        "and compact evidence before a language-model response is requested. It does not receive the full raw dataset."
    ))
    add_image(document, "04_sentiment.png", 5.8,
        "Figure 1. Sentiment dashboard with branch selector, customer-mood distribution, and branch comparison.")
    add_image(document, "10_sentiment_wordcloud.png", 5.8,
        "Figure 2. Review-language word cloud. Term size indicates frequency in selected customer feedback and should be interpreted with its review context.")

    add_heading(document, "5. Findings and CRM Insights")
    add_para(document, (
        f"The text baseline classified {sentiment['Positive']:,} records as positive, {sentiment['Negative']:,} as negative, and "
        f"{sentiment['Neutral']:,} as neutral. The large neutral group demonstrates why a lexicon output must not be treated as a "
        "validated CSAT or NPS score. It is used as a screening signal alongside observed rating, review volume, and manager inspection."
    ))
    add_table(document,
        ["Branch", "Reviews", "Avg. rating", "Text available", "CRM reading"],
        [[b["branch_name"], f"{int(b['review_count']):,}", b["average_rating"], f"{int(b['text_available_count']):,}",
          "Compare operational context before allocating response effort."] for b in branches],
        [1.12, 0.72, 0.85, 1.00, 2.81], font_size=9.5)
    add_para(document, "Table 1. Observed branch coverage and ratings. Source: project-generated branch summary from collected Google Maps exports.", italic=True, after=6)
    add_para(document, (
        f"The most frequent detected experience-area mentions are food quality ({aspect_counts['food_quality']:,}), ambience "
        f"({aspect_counts['ambience']:,}), service ({aspect_counts['service']:,}), price/value ({aspect_counts['price_value']:,}), "
        f"and waiting time ({aspect_counts['waiting_time']:,}). Frequency indicates where to look first, not the cause of a customer outcome."
    ))
    add_image(document, "02_branch_comparison_chart.png", 5.8,
        "Figure 3. Observed average rating and collected review count by branch. Volume differences should be considered when comparing branches.")

    add_heading(document, "6. From Insight to CRM Action")
    add_para(document, (
        "The CRM priority engine creates a case when recurring experience-area signals are combined with branch alert information. "
        "Every case includes its supporting mention count, an action owner, a status, and a weekly review instruction. This means "
        "a manager can use the dashboard to triage attention without claiming that a model has proven the root cause."
    ))
    banjara = [row for row in crm_cases if row["branch_id"] == "banjara_hills"]
    selected = sorted(banjara, key=lambda row: -int(row["mention_count"]))[:4]
    add_table(document,
        ["Experience area", "Mentions", "Priority", "CRM response"],
        [[row["issue"].replace("_", " ").title(), row["mention_count"], row["priority"].title(), row["recommended_action"]]
         for row in selected],
        [1.45, 0.70, 0.80, 3.55], font_size=9.4)
    add_para(document, "Table 2. Sample evidence-ranked Banjara Hills CRM cases. Source: deterministic project CRM case artifact.", italic=True, after=6)
    add_labeled_para(document, "Service recovery trigger. ", (
        "When a branch-month alert coincides with recurring service or wait-time language, review a sample of recent low-rating comments, "
        "check the operating period, assign a branch owner within one business day, and track the action to closure."
    ))
    add_labeled_para(document, "Proactive outreach rule. ", (
        "Where contact and platform policy permit, a customer-support lead should acknowledge strong dissatisfaction promptly, avoid making "
        "unverified claims, offer a suitable recovery route, and record the interaction in the CRM case history."
    ))
    add_labeled_para(document, "Closed-loop VoC routine. ", (
        "Capture the review, classify the issue, validate a sample, respond or improve the operating process, and monitor the same "
        "experience area in the next review cycle. This converts a social-review touchpoint into an accountable customer-relationship process."
    ))
    add_image(document, "07_anomaly_center.png", 5.8,
        "Figure 4. Anomaly Center page. Observed branch-month ratings are marked when the monitoring rules indicate manager attention.")

    add_heading(document, "7. AI Analyst: Controlled CRM Assistance")
    add_para(document, (
        "The AI Analyst supports a manager after deterministic analytics have prepared the evidence. The orchestration layer uses rule-based "
        "branch and period detection, then gathers compact JSON evidence such as branch summary, topic terms, aspect signals, alerts, and CRM "
        "cases. The CRM recommendation agent is instructed to ground every action in that evidence, label confidence based on real, derived, or "
        "predicted information, and avoid fabricated facts."
    ))
    add_labeled_para(document, "Example CRM question. ", "What actions should we take for Banjara Hills?")
    add_labeled_para(document, "Expected tool path. ", (
        "The orchestrator detects the branch and an action intent, calls the insight agent with compact evidence, and then calls the CRM "
        "recommendation agent. The output is an explanation and prioritized suggestions, not an autonomous customer-facing decision."
    ))
    add_image(document, "11_ai_scenario_banjara_full_page.png", 5.8,
        "Figure 5. AI Analyst scenario for Banjara Hills, showing the complete page, manager question, evidence-grounded answer, and full navigation.")

    add_heading(document, "8. Implementation Roadmap and Risk Mitigation")
    add_para(document, (
        "Implementation should begin as a controlled branch pilot, not an automatic decision system. The recommended process adds feedback fields "
        "to the CRM or service log, defines an owner for each high-priority issue, establishes review rules for alert months, and checks whether "
        "actions correspond with later observed feedback patterns."
    ))
    add_gantt_chart(document)
    add_para(document, "Figure 6. Twelve-week implementation roadmap for integrating feedback signals into the CRM operating rhythm.", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=6)
    add_table(document,
        ["Risk", "Mitigation"],
        [
            ["False or mixed sentiment", "Use sentiment as a triage cue; validate a sample of reviews before action."],
            ["Uneven branch volume", "Compare rates and supporting comments, not counts alone."],
            ["Unclear accountability", "Assign named owner, due date, case status, and weekly review."],
            ["Model drift or language change", "Recheck thresholds and manually review a labeled sample each quarter."],
            ["Over-automation", "Keep recovery and operational decisions under human approval."],
        ], [1.55, 4.95], font_size=9.8)
    add_para(document, "Table 3. CRM implementation risks and mitigation controls.", italic=True, after=6)
    add_labeled_para(document, "Proposed CRM scorecard. ", (
        "Track operational measures first: percentage of high-priority cases reviewed within one business day, case-closure time, "
        "repeat mentions by experience area, review volume, and rating trend. Add CSAT, repeat visit, loyalty, or lifetime-value metrics "
        "only when a consented CRM data source makes them measurable."
    ))

    add_heading(document, "9. Ethical, Legal, and Bias Considerations")
    add_para(document, (
        "Online reviews can contain personal information, informal language, sarcasm, mixed opinions, and writing in multiple languages or dialects. "
        "The project minimizes data fields used for analysis, excludes unnecessary reviewer metadata from analytical outputs, and keeps raw source files separate. "
        "Any future outreach must follow the source platform's policies and applicable privacy requirements."
    ))
    add_labeled_para(document, "Bias and interpretability. ", (
        "The current sentiment method is a transparent lexicon baseline and is explicitly marked derived and unvalidated. It can miss negation, sarcasm, "
        "context, and non-English nuance. A human review sample is required before the model influences high-impact customer treatment."
    ))
    add_labeled_para(document, "Transparency. ", (
        "The dashboard labels collected, derived, and predicted evidence separately. The AI Analyst is constrained to evidence bundles and should present a "
        "recommendation as an input to a manager, not an automated decision."
    ))

    add_heading(document, "10. Conclusion")
    add_para(document, (
        "This CRM Track C project demonstrates how Mandi @ 36 can turn large-scale unstructured feedback into a disciplined Voice of Customer workflow. "
        "The contribution is a practical linkage between customer language, branch context, alert signals, and accountable CRM response. The system is useful "
        "because it is auditable: review records remain the evidence base, derived NLP outputs are labeled, and a manager validates the action."
    ))
    add_callout(document, "Recommended next step.", (
        "Pilot the workflow for four weeks at Banjara Hills and Gachibowli. Review high-priority CRM cases weekly, record actions and outcomes, and use a manually "
        "labeled sample of feedback to calibrate the sentiment and aspect models before any wider operational reliance."
    ))

    add_heading(document, "Appendix A. Access, Reproducibility, and Data Provenance")
    add_para(document, "Repository: ", after=2)
    add_hyperlink(document.paragraphs[-1], "https://github.com/im45145v/Mandi360", "https://github.com/im45145v/Mandi360")
    add_para(document, "Live dashboard: ", after=2)
    add_hyperlink(document.paragraphs[-1], "https://mandi36.streamlit.app/", "https://mandi36.streamlit.app/")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    lead = paragraph.add_run("Google Maps review sources: ")
    set_run_font(lead, size=11, bold=True)
    add_hyperlink(paragraph, "Banjara Hills", "https://maps.app.goo.gl/Q3ThcUTcjc2UsFoS8")
    separator = paragraph.add_run(" | ")
    set_run_font(separator, size=11)
    add_hyperlink(paragraph, "Gachibowli", "https://maps.app.goo.gl/5rGRF4CdBzZYRULw8")
    separator = paragraph.add_run(" | ")
    set_run_font(separator, size=11)
    add_hyperlink(paragraph, "Jubilee Hills", "https://maps.app.goo.gl/DTt6HFX2THJhiixg8")
    add_para(document, (
        "Reproducible entry points: src/pipeline.py produces the analytical artifacts; app/Home.py starts the dashboard; "
        "src/agents/orchestrator.py routes evidence-grounded CRM questions."
    ))

    add_heading(document, "References")
    references = [
        "Grootendorst, M. (2022). BERTopic: Neural topic modeling with a class-based TF-IDF procedure. arXiv. https://arxiv.org/abs/2203.05794",
        "Kumar, V., & Reinartz, W. (2018). Customer relationship management: Concept, strategy, and tools. Springer.",
        "Liu, B. (2012). Sentiment analysis and opinion mining. Morgan & Claypool Publishers.",
        "Özdağoğlu, G., Kapuçugil Ikiz, A., & Celik, A. F. (2018). Topic modelling-based decision framework for analysing digital voice of the customer. Total Quality Management & Business Excellence, 29, 1545-1562. https://doi.org/10.1080/14783363.2016.1273106",
        "Pang, B., & Lee, L. (2008). Opinion mining and sentiment analysis. Foundations and Trends in Information Retrieval, 2(1-2), 1-135.",
        "Suresh, S., Rajan, G. T., & Gopinath, V. (2018). VoC-DL: Revisiting Voice of Customer using deep learning. Proceedings of AAAI, 32(1). https://doi.org/10.1609/aaai.v32i1.11408",
        "Winer, R. S. (2001). A framework for customer relationship management. California Management Review, 43(4), 89-105. https://doi.org/10.2307/41166102",
        "Zhang, J. (2019). What's yours is mine: Exploring customer voice on Airbnb using text-mining approaches. Journal of Consumer Marketing, 36(5), 655-665. https://doi.org/10.1108/JCM-02-2018-2581",
    ]
    for reference in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(reference)
        set_run_font(run, size=12)

    document.save(OUTPUT)


def build_logbook() -> None:
    document = Document()
    setup_document(document, "Mandi @ 36 | CRM Prompt Logbook")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("CRM Track C Prompt Logbook")
    set_run_font(run, size=22, bold=True, color=INK)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Mandi @ 36: AI-Driven Sentiment Analysis of Customer Feedback")
    set_run_font(run, size=13, italic=True, color=BLUE)
    add_para(document, f"Prepared by: {NAME} | Roll Number: {ROLL}", align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_callout(document, "Purpose.", "This logbook records the practical prompts used or reconstructed from the Mandi @ 36 workflow. AI assisted with code scaffolding, interpretation, dashboard communication, and writing. It did not create raw customer reviews or replace deterministic analytics.")
    add_heading(document, "Prompt Record")
    add_para(document, (
        "The prompts below are written in a reproducible form so that a reviewer can understand what information was shared, what output was requested, "
        "and how the result was checked. Tools such as Copilot, Claude, Codex, and ChatGPT were used selectively according to the task. Any generated code or "
        "narrative was reviewed against the local files and project outputs before inclusion."
    ))
    records = [
        ("1", "Codex\nSchema inspection", "I have Apify Google Maps scraper JSON exports. Inspect the actual schema without changing any raw file. Identify the review-list field and map location, review text, rating, date, reviewer name, and source URL. Report missing or ambiguous fields before proposing code.", "Produced a field-mapping checklist and schema-validation approach. Used to keep raw files intact and document the available evidence."),
        ("2", "GitHub Copilot\nJSON preparation", "Convert the Google Maps scraper JSON into a usable analytical table in Python. Write a defensive loader that accepts nested review arrays, keeps branch name and source URL, normalizes rating and date, preserves raw text, and writes a separate cleaned output. Do not invent missing values.", "Provided a code scaffold for normalized review records. The implementation was adapted after checking the real export structure and saved outputs separately from raw data."),
        ("3", "Claude\nText cleaning", "Suggest a privacy-minimizing text-cleaning sequence for public restaurant reviews. Preserve the original review text, create a clean analysis field, and list limitations for emojis, mixed languages, sarcasm, and missing text.", "Produced cleaning steps and limitations. Used to explain preparation choices and the requirement for human interpretation."),
        ("4", "Codex\nSentiment baseline", "Create a transparent sentiment-analysis baseline for restaurant review text using interpretable positive and negative lexicons. Return a label and score, keep the logic separate from any LLM call, and state why it must not be presented as CSAT or NPS.", "Generated a deterministic baseline design. It informed the derived sentiment labels and the report's caution that model labels are screening signals."),
        ("5", "ChatGPT\nTopic and aspect interpretation", "Explain, in plain language, how keyword aspects and TF-IDF plus NMF topic modeling can identify recurring restaurant experience areas such as food, service, ambience, waiting time, and value. Include limits and manager-validation steps.", "Created a plain-language explanation. Used to make the methodology and finding sections understandable without overstating causation."),
        ("6", "GitHub Copilot\nDashboard design", "For a Streamlit restaurant Voice of Customer dashboard, propose 5 to 7 visuals that help a branch manager move from review text to action. Include sentiment distribution, branch comparison, word cloud, trend, complaint themes, anomaly flags, and action cases.", "Suggested dashboard sections and visual roles. The final application shows selected visuals and preserves the distinction between collected, derived, and predicted evidence."),
        ("7", "Project AI Analyst\nCRM action query", "What actions should we take for Banjara Hills? Use only the supplied compact JSON evidence for branch summary, recurring aspects, topics, alerts, and CRM cases. State evidence, confidence, recommended owner, next action, and what a manager should validate.", "Returned an evidence-grounded action narrative. A full page screenshot and example output are included in the report; the manager remains the final decision-maker."),
        ("8", "Claude\nEthics review", "Review this public-review sentiment workflow for privacy, bias, transparency, and over-automation risks. Give practical controls for a restaurant CRM process, including consent, retention, human review, and treatment of sarcasm or language variation.", "Produced a risk-and-control checklist. Used to strengthen the ethical, legal, and bias considerations section."),
        ("9", "ChatGPT\nCRM roadmap", "Propose a 12-week phased implementation roadmap for a restaurant feedback-to-action workflow. Include pilot, ownership, response SLA, staff training, monitoring, calibration, risks, and mitigation. Mark recommendations clearly rather than presenting them as completed operations.", "Provided a planning structure. It was converted into the Gantt chart, CRM scorecard, and risk table after adapting it to the project scope."),
        ("10", "Codex\nReport QA", "Check the report against this brief: 10 to 12 pages, Times New Roman 12 pt body text, 1.5 spacing, 1-inch margins, justified paragraphs, APA references, clickable deployment and repository links, Google Maps sources, visuals, and no unsupported claims. Identify only evidence-backed corrections.", "Produced a quality-assurance checklist. The final report was rendered and visually checked, with figures, links, citations, and claims reviewed before submission."),
    ]
    add_table(document,
        ["No.", "Tool and task", "Prompt text", "Output, project use, and reflection"],
        [[*record] for record in records],
        [0.36, 0.92, 3.15, 2.07], font_size=9.2, left_aligned_cols={1, 2, 3})
    add_heading(document, "Verification and Reflection")
    add_para(document, (
        "All computed numerical results in the CRM report are sourced from the local pipeline artifacts, including dataset summary, branch summary, NLP outputs, anomaly outputs, topic model outputs, and CRM cases. AI assistance was used to draft explanations and structure, but it was not used to manufacture review records, metrics, or customer outcomes. The strongest practice is to retain evidence links, label derived results, and require human validation before customer-facing action."
    ))
    add_heading(document, "Technical Evidence Used")
    add_labeled_para(document, "Pipeline. ", "src/pipeline.py reads raw exports, normalizes data, runs analytics, and writes result artifacts without modifying raw review files.")
    add_labeled_para(document, "Sentiment and aspects. ", "src/analytics/nlp.py implements a transparent lexicon baseline and keyword aspect extraction, each marked derived and unvalidated.")
    add_labeled_para(document, "CRM cases. ", "src/analytics/crm.py ranks recurring experience-area signals and attaches a suggested owner, action, status, and review cycle.")
    add_labeled_para(document, "AI routing. ", "src/agents/orchestrator.py detects branch, comparison, period, and action intent before passing compact evidence to the CRM recommendation agent.")
    document.save(LOGBOOK_OUTPUT)


if __name__ == "__main__":
    build_report()
    build_logbook()
